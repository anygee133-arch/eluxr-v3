#!/usr/bin/env python3
"""
Generate ELUXR-AI-Prompts.docx — a human-readable Word document
containing all AI prompts used in the ELUXR v3 content pipeline.

Run:  python3 /Users/andrewgershan/Projects/eluxr-v3/docs/generate_prompts_docx.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "ELUXR-AI-Prompts.docx")

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──────────────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x22, 0x22, 0x22)
style.paragraph_format.space_after = Pt(6)

# Heading styles
for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    hs.font.name = "Calibri"

# Custom "Prompt Text" style — the body of each prompt
prompt_style = doc.styles.add_style("PromptText", WD_STYLE_TYPE.PARAGRAPH)
prompt_style.font.name = "Consolas"
prompt_style.font.size = Pt(10)
prompt_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
prompt_style.paragraph_format.space_before = Pt(4)
prompt_style.paragraph_format.space_after = Pt(2)
prompt_style.paragraph_format.left_indent = Cm(0.5)

# "Meta" style for workflow/node/model info
meta_style = doc.styles.add_style("Meta", WD_STYLE_TYPE.PARAGRAPH)
meta_style.font.name = "Calibri"
meta_style.font.size = Pt(9)
meta_style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
meta_style.font.italic = True
meta_style.paragraph_format.space_after = Pt(2)

# "Placeholder" style for the placeholder legend
placeholder_style = doc.styles.add_style("Placeholder", WD_STYLE_TYPE.PARAGRAPH)
placeholder_style.font.name = "Calibri"
placeholder_style.font.size = Pt(9)
placeholder_style.font.color.rgb = RGBColor(0x88, 0x44, 0x00)
placeholder_style.paragraph_format.space_after = Pt(2)
placeholder_style.paragraph_format.left_indent = Cm(0.5)

# ── Helpers ─────────────────────────────────────────────────────────
def add_prompt_block(lines):
    """Add prompt text, splitting on newlines, each line in PromptText style."""
    for line in lines:
        doc.add_paragraph(line, style="PromptText")

def add_meta(text):
    doc.add_paragraph(text, style="Meta")

def add_placeholders(items):
    """items: list of (placeholder, description) tuples."""
    doc.add_paragraph("Placeholders used:", style="Meta")
    for ph, desc in items:
        doc.add_paragraph(f"  {ph}  --  {desc}", style="Placeholder")

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("_" * 80)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)


# ═══════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════
title = doc.add_heading("ELUXR v3 -- AI Prompts for Content Generation", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "Edit these prompts to change what the AI generates at each step.\n"
    "Dynamic values are shown as {PLACEHOLDERS}."
)
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()  # spacer

intro = doc.add_paragraph(
    "This document contains the exact prompts sent to the AI models "
    "(Claude, GPT-4o, GPT-4.1-mini) at each stage of the ELUXR content pipeline. "
    "Each prompt section shows:\n\n"
    "  1. What it generates\n"
    "  2. Which n8n workflow and node it lives in\n"
    "  3. Which AI model processes it\n"
    "  4. The full prompt text with {PLACEHOLDERS} for dynamic data\n"
    "  5. What each placeholder represents\n\n"
    "To change the AI's behavior, edit the prompt text and give it back "
    "to be re-implemented in the corresponding n8n workflow node."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
#  PROMPT 1: Product Extraction (Step 1)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("1. Step 1: Product Extraction", level=1)
doc.add_heading("Extracts products from a brand's website", level=3)
add_meta("Workflow: 01-ICP-Analyzer (Bin0AjccOtr2etgH)")
add_meta("Node: Prepare Extract Products")
add_meta("AI Model: Claude Sonnet 4 (claude-sonnet-4-20250514)")
doc.add_paragraph()

add_prompt_block([
    "SYSTEM MESSAGE:",
    "",
    "You are a product extractor for luxury brand websites. Extract as many real products as possible from the page content. Products can be found in navigation menus, category links, collection listings, product grids, or any other part of the page.",
    "",
    "---",
    "",
    "USER MESSAGE:",
    "",
    "Extract ALL individual products from this website content. Look everywhere: navigation menus, category links, collection pages, product listings, image alt text, breadcrumbs, and any product references. We need UP TO 30 unique products.",
    "",
    "{COMBINED_WEBSITE_CONTENT}",
    "",
    "Return a JSON array. Each product needs:",
    "- name: string (specific product name like \"Tiffany T Wire Bracelet\" or \"Rolling Diamonds Eternity Ring\", NOT generic categories like \"Bracelets\")",
    "- description: string (brief, max 100 chars -- infer from context if needed)",
    "- price: string or null",
    "- features: string[] (max 5 -- infer from product name/category)",
    "- image_url: string or null (only if an actual image URL appears in the content)",
    "- image_urls: string[] (only real URLs from the content)",
    "- product_url: string or null (must be a URL that actually appears in the page content, absolute URL starting with https://)",
    "- category: string (e.g. Rings, Necklaces, Bracelets, Earrings, Watches)",
    "",
    "RULES:",
    "1. Extract specific NAMED products, not generic categories. \"Tiffany T Wire Bracelet\" is good. \"Bracelets\" is not.",
    "2. Look in navigation menus -- collection names like \"Tiffany T\", \"HardWear\", \"Elsa Peretti\" are product lines. Extract individual products from them.",
    "3. If you see a collection name but no individual products, create entries for the collection's likely products (e.g. \"Tiffany T Smile Pendant\", \"Tiffany T Wire Ring\").",
    "4. product_url must be from the actual page content. Do not guess URLs.",
    "5. Try to reach at least 20 products. Use product names, collection references, and category links to identify as many as possible.",
    "6. Base domain: {BUSINESS_URL}",
    "",
    "Return ONLY a valid JSON array. No markdown.",
])

doc.add_paragraph()
add_placeholders([
    ("{COMBINED_WEBSITE_CONTENT}", "Scraped content from up to 4 pages: homepage + 3 product listing pages (via Jina Reader), truncated to ~70K chars total."),
    ("{BUSINESS_URL}", "The brand's website URL entered by the user (e.g. heinzmayer.de, tiffany.com)."),
])

add_separator()


# ═══════════════════════════════════════════════════════════════════
#  PROMPT 2: ICP Synthesis (Step 1)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("2. Step 1: ICP Synthesis (Ideal Customer Profile)", level=1)
doc.add_heading("Combines 3 Perplexity research rounds + products into a full ICP", level=3)
add_meta("Workflow: 01-ICP-Analyzer (Bin0AjccOtr2etgH)")
add_meta("Node: Prepare Synthesize ICP")
add_meta("AI Model: Claude Sonnet 4 (claude-sonnet-4-20250514)")
doc.add_paragraph()

add_prompt_block([
    "You are a world-class marketing strategist. Based on the following research, create a comprehensive Ideal Customer Profile (ICP) for this business.",
    "",
    "BUSINESS: {BUSINESS_URL}",
    "INDUSTRY: {INDUSTRY}",
    "",
    "WEBSITE CONTENT SUMMARY:",
    "{SITE_CONTENT}",
    "",
    "PRODUCTS/SERVICES OFFERED:",
    "{PRODUCTS_LIST}",
    "",
    "RESEARCH 1 - INDUSTRY LANDSCAPE & COMPETITORS:",
    "{RESEARCH_1}",
    "",
    "RESEARCH 2 - AUDIENCE PAIN POINTS & DESIRES:",
    "{RESEARCH_2}",
    "",
    "RESEARCH 3 - CONTENT GAPS & OPPORTUNITIES:",
    "{RESEARCH_3}",
    "",
    "Synthesize ALL of the above into a professional ICP. Return a JSON object with EXACTLY these keys:",
    "{",
    "  \"icp_summary\": \"2-3 paragraph executive summary of the ideal customer\",",
    "  \"demographics\": { \"age_range\": \"\", \"gender\": \"\", \"income_level\": \"\", \"education\": \"\", \"location\": \"\", \"job_titles\": [] },",
    "  \"psychographics\": { \"values\": [], \"interests\": [], \"lifestyle\": \"\" },",
    "  \"pain_points\": [\"array of specific pain points\"],",
    "  \"desires\": [\"array of specific desires and aspirations\"],",
    "  \"objections\": [\"array of common objections to purchasing\"],",
    "  \"buying_triggers\": [\"array of triggers that prompt purchasing\"],",
    "  \"content_pillars\": [{ \"name\": \"\", \"description\": \"\", \"example_topics\": [] }],",
    "  \"content_preferences\": { \"platforms\": {}, \"content_types\": [], \"best_posting_times\": {} },",
    "  \"competitors\": [{ \"name\": \"\", \"url\": \"\", \"strengths\": [], \"weaknesses\": [] }],",
    "  \"content_opportunities\": [\"array of underserved content topics\"],",
    "  \"recommended_hashtags\": { \"linkedin\": [], \"instagram\": [], \"x\": [], \"youtube\": [] }",
    "}",
    "",
    "Fill in ALL fields with real, researched data. Return ONLY valid JSON. No markdown, no explanation.",
])

doc.add_paragraph()
add_placeholders([
    ("{BUSINESS_URL}", "The brand's website URL."),
    ("{INDUSTRY}", "Industry entered by the user (e.g. 'luxury jewelry')."),
    ("{SITE_CONTENT}", "First ~3000 chars of the scraped website content."),
    ("{PRODUCTS_LIST}", "Formatted list: '- Product Name ($Price): Description' for each extracted product."),
    ("{RESEARCH_1}", "Perplexity research output: industry landscape, competitors, content strategies."),
    ("{RESEARCH_2}", "Perplexity research output: audience pain points, desires, objections, buying triggers."),
    ("{RESEARCH_3}", "Perplexity research output: content gaps, high-engagement formats, trending hashtags."),
])

add_separator()


# ═══════════════════════════════════════════════════════════════════
#  PROMPT 3: Theme / Show Generation (Step 3)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("3. Step 3: Theme / Show Generation (Netflix Model)", level=1)
doc.add_heading("Creates a monthly show name, 4 weekly seasons, and daily product assignments", level=3)
add_meta("Workflow: 02-Theme-Generator (hQeqQ0r6Ahop3YOI)")
add_meta("Node: Prepare Theme Prompt")
add_meta("AI Model: Claude Sonnet 4 (claude-sonnet-4-20250514)")
doc.add_paragraph()

add_prompt_block([
    "You are an expert social media content strategist using the NETFLIX MODEL for content planning.",
    "",
    "You are creating a monthly content campaign for a business. Generate a complete Netflix-model campaign structure.",
    "",
    "## BUSINESS CONTEXT",
    "- Industry: {INDUSTRY}",
    "- Month: {MONTH}",
    "- ICP Summary: {ICP_SUMMARY}",
    "- Content Pillars: {CONTENT_PILLARS}",
    "- Pain Points: {PAIN_POINTS}",
    "- Desires: {DESIRES}",
    "",
    "## PRODUCTS/SERVICES ({PRODUCT_COUNT} total)",
    "{PRODUCT_LIST}",
    "",
    "## NETFLIX MODEL RULES",
    "",
    "1. SHOW NAME: Create a catchy, branded show name for this month's campaign (e.g., \"The Glow-Up Series\", \"Build Different\", \"The Founders Table\")",
    "",
    "2. 4 SEASONS (1 per week), each with a progressive arc:",
    "   - Season 1 (Week 1): \"intro\" -- Introduce the brand, products, and value proposition. Hook the audience.",
    "   - Season 2 (Week 2): \"deepen\" -- Go deeper into product benefits, customer stories, expertise. Build trust.",
    "   - Season 3 (Week 3): \"challenge\" -- Address objections, myths, comparisons. Overcome barriers.",
    "   - Season 4 (Week 4): \"celebrate\" -- Success stories, community, future vision. Create loyalty.",
    "",
    "3. DAILY ASSIGNMENTS: Each day gets exactly 1 product + 1 inspirational topic",
    "   - 7 days per season = 28 days total",
    "   - Products should cluster by theme coherence (similar products in same week when possible)",
    "   - [Note: If more than 30 products, the top 30 most ICP-relevant are selected. If fewer products than days, products are distributed evenly with repeats.]",
    "",
    "4. WEEKLY HOOKS: Last day of each week teases the next season. Daily posts stand alone within their week.",
    "",
    "## REQUIRED OUTPUT FORMAT",
    "",
    "Return a JSON object with this EXACT structure:",
    "{",
    "  \"show_name\": \"The [Catchy Name] Series\",",
    "  \"show_concept\": \"Brief description of the month's overarching narrative (1-2 sentences)\",",
    "  \"seasons\": [",
    "    {",
    "      \"week_number\": 1,",
    "      \"theme_name\": \"Season 1: [Creative Name]\",",
    "      \"inspirational_theme\": \"Introduction/Discovery theme description\",",
    "      \"season_arc\": \"intro\",",
    "      \"theme_description\": \"What this week covers and why it matters\",",
    "      \"hook\": \"Weekly hook teasing next season (for last post of week)\",",
    "      \"days\": [",
    "        {",
    "          \"day_number\": 1,",
    "          \"product_id\": \"uuid-of-product-or-null\",",
    "          \"product_name\": \"Product Name\",",
    "          \"topic\": \"Inspirational daily topic tied to product + theme\",",
    "          \"content_angle\": \"How to present this product through the theme lens\"",
    "        }",
    "      ]",
    "    }",
    "  ]",
    "}",
    "",
    "Each season MUST have exactly 7 days. Each day MUST have a product_id (use actual UUIDs from the product list above), product_name, topic, and content_angle.",
    "",
    "Return ONLY valid JSON. No markdown fences. No explanation.",
])

doc.add_paragraph()
add_placeholders([
    ("{INDUSTRY}", "Industry from user input."),
    ("{MONTH}", "Target content month (e.g. '2026-04')."),
    ("{ICP_SUMMARY}", "ICP summary from the ICP table (or 'General audience' if missing)."),
    ("{CONTENT_PILLARS}", "JSON array of content pillar objects from the ICP."),
    ("{PAIN_POINTS}", "JSON array of pain point strings from the ICP."),
    ("{DESIRES}", "JSON array of desires from the ICP."),
    ("{PRODUCT_COUNT}", "Number of extracted products."),
    ("{PRODUCT_LIST}", "Numbered list: '1. **Product Name** (ID: uuid) / Description / Price / Category / Features' for each product."),
])

add_separator()


# ═══════════════════════════════════════════════════════════════════
#  PROMPT 4: Topic Generation (Step 3)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("4. Step 3: Daily Topic Generation (7-Day Storytelling)", level=1)
doc.add_heading("Generates 7 daily content topics for one week, following the storytelling framework", level=3)
add_meta("Workflow: 15-Generate-Topics (qTDGBLtfBqyb1Vm1)")
add_meta("Node: Prepare Claude Prompt")
add_meta("AI Model: Claude Sonnet 4 (claude-sonnet-4-20250514)")
doc.add_paragraph()

add_prompt_block([
    "SYSTEM MESSAGE:",
    "",
    "You are a luxury brand social media content strategist specializing in the 7-Day Storytelling Framework. Generate daily topics that follow the weekly narrative arc.",
    "",
    "---",
    "",
    "USER MESSAGE:",
    "",
    "Generate 7 daily social media content topics for WEEK {WEEK_NUMBER} of a monthly campaign.",
    "",
    "BUSINESS CONTEXT:",
    "- Industry: {INDUSTRY}",
    "- Target Audience: {TARGET_AUDIENCE}",
    "- Brand Voice: {BRAND_VOICE}",
    "- Content Pillars: {CONTENT_PILLARS}",
    "- Pain Points: {PAIN_POINTS}",
    "",
    "PRODUCT ASSIGNMENTS (MANDATORY - each day MUST use its assigned product):",
    "{PRODUCT_LIST}",
    "",
    "7-DAY STORYTELLING FRAMEWORK (MANDATORY - each day MUST follow its assigned theme):",
    "Day 1 (Monday): CURIOSITY HOOK -- Stop the scroll with intrigue and visual mystery.",
    "Day 2 (Tuesday): CRAFTSMANSHIP -- Establish luxury credibility and authority.",
    "Day 3 (Wednesday): MEANING & SYMBOLISM -- Build emotional connection.",
    "Day 4 (Thursday): TRANSFORMATION MOMENT -- Show the impact when the product is worn/used.",
    "Day 5 (Friday): HIDDEN DETAIL / SURPRISE -- Encourage shares through discovery.",
    "Day 6 (Saturday): LIFESTYLE ASPIRATION -- Inspire desire and aspiration.",
    "Day 7 (Sunday): LEGACY & TIMELESSNESS -- Build brand prestige.",
    "",
    "IMPORTANT RULES:",
    "- Each day MUST use the EXACT product assigned above (Day 1 product for day 1, Day 2 product for day 2, etc.)",
    "- Do NOT reassign or swap products between days",
    "- The product_name in your response MUST exactly match the assigned product name",
    "",
    "{ADDITIONAL_STORYTELLING_CONTEXT}",
    "",
    "For each of the 7 days provide:",
    "1. topic: A compelling topic title that reflects the day's storytelling theme (max 60 chars)",
    "2. description: Brief description of the post angle tied to the theme (max 150 chars)",
    "3. content_angle: Specific approach for this theme",
    "4. product_name: The EXACT product name assigned to this day (from the list above)",
    "5. storytelling_theme: The theme name",
    "",
    "Return ONLY a valid JSON array of 7 objects with keys: day_number (1-7), topic, description, content_angle, product_name, storytelling_theme. No markdown.",
])

doc.add_paragraph()
add_placeholders([
    ("{WEEK_NUMBER}", "Which week of the month (1-4)."),
    ("{INDUSTRY}", "Industry from the ICP record."),
    ("{TARGET_AUDIENCE}", "target_audience or demographics from the ICP."),
    ("{BRAND_VOICE}", "brand_voice or tone from the ICP."),
    ("{CONTENT_PILLARS}", "JSON array of content pillar objects from the ICP."),
    ("{PAIN_POINTS}", "JSON array of pain point strings from the ICP."),
    ("{PRODUCT_LIST}", "One line per day: 'Day 1: Product Name (Category) -- Description'."),
    ("{ADDITIONAL_STORYTELLING_CONTEXT}", "Optional: extra storytelling framework text from frontend (if provided)."),
])

add_separator()


# ═══════════════════════════════════════════════════════════════════
#  PROMPT 5: Content Generation / megaPrompt (Step 4)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("5. Step 4: Platform Content Generation (the megaPrompt)", level=1)
doc.add_heading("Generates 7 days of multi-platform social media posts in one call", level=3)
add_meta("Workflow: 04-Content-Studio (TreszUaJqlykCrMi)")
add_meta("Node: Prepare Daily Content Loop")
add_meta("AI Model: Claude Sonnet 4 (claude-sonnet-4-20250514)")
doc.add_paragraph()

add_prompt_block([
    "You are a luxury brand content strategist. Generate social media content for ALL 7 days following the Viral Storytelling Framework.",
    "",
    "=== STORYTELLING FRAMEWORK ===",
    "Weekly rhythm: Curiosity > Craftsmanship > Emotion > Transformation > Surprise > Aspiration > Legacy",
    "",
    "=== CONTEXT ===",
    "Show: \"{SHOW_NAME}\"",
    "Season {WEEK_NUMBER}: \"{THEME_NAME}\" -- {INSPIRATIONAL_THEME}",
    "ICP: {ICP_SUMMARY}",
    "Platforms: {PLATFORMS}",
    "",
    "=== VIDEO NOTE ===",
    "For Instagram and YouTube, videos will use the ACTUAL PRODUCT IMAGE. Scripts should describe content anchored to real product photos.",
    "",
    "=== 7-DAY PLAN ===",
    "",
    "[For each of the 7 days, the following block is included:]",
    "",
    "Day {DAY_NUMBER} ({DAY_NAME} - {STORYTELLING_THEME}):",
    "  Goal: {GOAL}",
    "  Concept: {CONCEPT}",
    "  Angle: {ANGLE}",
    "  Hooks: \"{HOOK_1}\", \"{HOOK_2}\", \"{HOOK_3}\"",
    "  Driver: {DRIVER}",
    "  Product: {PRODUCT_NAME} -- {PRODUCT_DESCRIPTION}",
    "  Topic: \"{TOPIC}\"",
    "  Angle: \"{CONTENT_ANGLE}\"",
    "",
    "[The 7-day storytelling themes are hardcoded as:]",
    "",
    "Day 1 (Monday): Curiosity Hook",
    "  Goal: Stop the scroll with intrigue and visual mystery",
    "  Concept: Introduce the piece through partial reveal, light reflection, or extreme macro details that spark curiosity",
    "  Hooks: \"The detail most people never notice...\", \"A piece that changes everything.\", \"Look closer.\"",
    "  Driver: Curiosity increases watch time and replays",
    "",
    "Day 2 (Tuesday): Craftsmanship",
    "  Goal: Establish luxury credibility and authority",
    "  Concept: Highlight craftsmanship behind the piece",
    "  Hooks: \"From raw material to wearable art.\", \"Precision in every millimeter.\", \"Crafted, not manufactured.\"",
    "  Driver: Craftsmanship content triggers fascination and trust",
    "",
    "Day 3 (Wednesday): Meaning & Symbolism",
    "  Goal: Build emotional connection",
    "  Concept: Position jewelry as a symbol of moments, relationships, or personal identity",
    "  Hooks: \"Some pieces mark a moment forever.\", \"Jewelry is never just jewelry.\", \"Every piece tells a story.\"",
    "  Driver: Emotion increases saves and comments",
    "",
    "Day 4 (Thursday): Transformation Moment",
    "  Goal: Show the impact when the jewelry is worn",
    "  Concept: The transformation from object to presence",
    "  Hooks: \"The moment everything changes.\", \"From detail to statement.\", \"Presence begins here.\"",
    "  Driver: Transformation visuals drive shares and engagement",
    "",
    "Day 5 (Friday): Hidden Detail / Surprise",
    "  Goal: Encourage shares through discovery",
    "  Concept: Reveal something unexpected about the piece, gemstone, or design",
    "  Hooks: \"Most people never notice this.\", \"The secret behind the sparkle.\", \"Why this cut catches light differently.\"",
    "  Driver: Surprise drives sharing behavior",
    "",
    "Day 6 (Saturday): Lifestyle Aspiration",
    "  Goal: Inspire desire and aspiration",
    "  Concept: Place jewelry in aspirational environments",
    "  Hooks: \"Luxury is a feeling.\", \"Some moments deserve more.\", \"Where elegance begins.\"",
    "  Driver: Aspirational content fuels desire",
    "",
    "Day 7 (Sunday): Legacy & Timelessness",
    "  Goal: Build brand prestige and long-term emotional value",
    "  Concept: Jewelry as heirloom, tradition, and timeless symbol",
    "  Hooks: \"Some pieces outlive trends.\", \"Designed to last generations.\", \"What will your story be?\"",
    "  Driver: Legacy messaging strengthens brand authority",
    "",
    "=== PLATFORM RULES ===",
    "- LinkedIn: 1300-1800 chars, thought leadership tone, line breaks between paragraphs, 3-5 professional hashtags, product CTA",
    "- Instagram: 150-300 char caption, 20-30 hashtags in separate first_comment field, emoji allowed, product CTA. Video will be generated from product image.",
    "- X: 5-7 tweet thread (each under 280 chars, numbered), punchy and conversational, product CTA in last tweet",
    "- YouTube: Detailed 30-60 second video script for YouTube Short. Include: title, hook (first 3 sec), 4 scene breakdowns with timing/visual/on_screen_text/narration, music_mood, thumbnail_concept, cta. Video generated from actual product image.",
    "- Facebook: 300-600 chars, conversational and engaging tone, emoji allowed, 3-5 hashtags, product CTA",
    "- TikTok: 30-60 second video script with hook in first 2 sec, trending format, casual tone, 3-5 hashtags, product CTA",
    "[Note: Only the platforms selected by the user are included.]",
    "",
    "=== OUTPUT FORMAT ===",
    "Return JSON:",
    "{",
    "  \"days\": [",
    "    {",
    "      \"day_number\": 1,",
    "      \"storytelling_theme\": \"Curiosity Hook\",",
    "      \"posts\": {",
    "        \"linkedin\": { \"text\": \"...\", \"hashtags\": [...], \"cta\": \"...\" },",
    "        \"instagram\": { \"text\": \"...\", \"hashtags\": [...], \"first_comment\": \"...\", \"cta\": \"...\" },",
    "        \"x\": { \"text\": \"...\", \"hashtags\": [...], \"cta\": \"...\" },",
    "        \"youtube\": { \"title\": \"...\", \"hook\": \"...\", \"scenes\": [...], \"music_mood\": \"...\", \"thumbnail_concept\": \"...\", \"cta\": \"...\", \"hashtags\": [...] },",
    "        \"facebook\": { \"text\": \"...\", \"hashtags\": [...], \"cta\": \"...\" },",
    "        \"tiktok\": { \"title\": \"...\", \"hook\": \"...\", \"scenes\": [...], \"music_mood\": \"...\", \"cta\": \"...\", \"hashtags\": [...] }",
    "      },",
    "      \"image_prompt\": \"Detailed AI image prompt for this day (max 500 chars)\"",
    "    },",
    "    ... (all 7 days)",
    "  ]",
    "}",
    "Include ALL platforms for EVERY day. Do not skip any.",
])

doc.add_paragraph()
add_placeholders([
    ("{SHOW_NAME}", "The Netflix-model show name from the campaign (e.g. 'The Brilliance Chronicles')."),
    ("{WEEK_NUMBER}", "Current week number (1-4)."),
    ("{THEME_NAME}", "The season/theme name for this week."),
    ("{INSPIRATIONAL_THEME}", "The inspirational theme description for this season."),
    ("{ICP_SUMMARY}", "Truncated ICP summary (first 500 chars)."),
    ("{PLATFORMS}", "Comma-separated list of target platforms selected by the user."),
    ("{DAY_NUMBER}", "1-7 for each day."),
    ("{DAY_NAME}", "Monday through Sunday."),
    ("{STORYTELLING_THEME}", "The hardcoded storytelling theme for this day (see list above)."),
    ("{GOAL}, {CONCEPT}, {ANGLE}, {HOOKS}, {DRIVER}", "Hardcoded storytelling framework details (see the 7 themes above)."),
    ("{PRODUCT_NAME}", "The product assigned to this day from the weekly topics."),
    ("{PRODUCT_DESCRIPTION}", "The product's description from the database."),
    ("{TOPIC}", "The approved topic for this day from the weekly topics."),
    ("{CONTENT_ANGLE}", "The approved content angle for this day from the weekly topics."),
])

add_separator()


# ═══════════════════════════════════════════════════════════════════
#  PROMPT 6: Image Analysis (Step 4)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("6. Step 4: Product Image Analysis", level=1)
doc.add_heading("Analyzes a product photo to extract visual details for image generation", level=3)
add_meta("Workflow: Social Calendar - Image Generator (Yh5DEtB1lR9lkbzo)")
add_meta("Node: Analyze Product Image1")
add_meta("AI Model: GPT-4o (vision)")
doc.add_paragraph()

add_prompt_block([
    "[The product image at {PRODUCT_IMAGE_URL} is sent to GPT-4o's vision API along with this text prompt:]",
    "",
    "Analyze the uploaded image and describe it in cinematic, production-ready terms.",
    "",
    "Return JSON with these fields:",
    "",
    "{",
    "  \"main_subject\": \"primary focus (e.g., man running, perfume bottle, luxury car, smiling woman)\",",
    "  \"subject_position\": \"center / left / right / close-up / wide / top-down / etc.\",",
    "  \"environment\": \"describe background or setting (e.g., beach at sunset, modern office, urban street, green forest)\",",
    "  \"lighting\": \"type and quality of light (e.g., warm golden hour, moody shadows, bright studio light)\",",
    "  \"color_palette\": \"dominant tones (e.g., beige and gold, blue and white, neon pink and cyan)\",",
    "  \"style\": \"visual or artistic style (e.g., cinematic realism, minimalistic fashion, tech aesthetic, natural documentary)\",",
    "  \"emotion_tone\": \"overall mood (e.g., calm, excitement, luxury, mystery, inspiration)\",",
    "  \"action_or_motion\": \"visible or implied movement (e.g., walking, wind, pouring, camera pan)\",",
    "  \"camera_perspective\": \"type of shot (macro, medium, wide, tracking, overhead, dolly)\",",
    "  \"potential_genre\": \"product_commercial / fashion / lifestyle / travel / tech / cinematic / nature / art\",",
    "  \"potential_vibe\": \"short creative tone summary (e.g., modern elegance, energetic motion, serene minimalism)\",",
    "  \"contextual_elements\": \"secondary visual hints (props, reflections, accessories, surrounding textures, lighting accents)\",",
    "  \"product_exact_colors\": \"e.g., matte black with rose gold trim\",",
    "  \"product_shape_silhouette\": \"e.g., rectangular box, curved bottle, circular face\",",
    "  \"product_text_or_logo\": \"any visible brand name, text, or logo on product\",",
    "  \"product_material_finish\": \"e.g., glossy glass, matte leather, brushed metal\",",
    "  \"product_unique_identifiers\": \"any distinctive marks, patterns, or design details\",",
    "  \"stone_orientation_precise\": \"describe exact cardinal orientation of stone on band -- e.g., heart point facing toward finger base (downward), lobes facing away from finger (upward), heart cleavage centered at 12 o'clock\",",
    "  \"stone_rotation_lock\": \"single sentence locking orientation -- e.g., the heart diamond sits with both lobes at the top and the point at the bottom, centered upright at 12 o'clock on the band, NOT rotated sideways\"",
    "}",
    "",
    "Guidelines:",
    "Keep phrasing professional and descriptive, not poetic.",
    "Always fill all fields (infer if needed).",
    "Return JSON only, no comments or text outside braces.",
])

doc.add_paragraph()
add_placeholders([
    ("{PRODUCT_IMAGE_URL}", "The URL of the product's image from the products table. Sent as the image input to GPT-4o vision."),
])

add_separator()


# ═══════════════════════════════════════════════════════════════════
#  PROMPT 7: Image Scene Prompt (Step 4)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("7. Step 4: Image Scene Prompt Generation", level=1)
doc.add_heading("Takes the image analysis + topic and writes a photorealistic image generation prompt", level=3)
add_meta("Workflow: Social Calendar - Image Generator (Yh5DEtB1lR9lkbzo)")
add_meta("Node: Image Scene Prompt1")
add_meta("AI Model: GPT-4.1-mini (via OpenAI Chat Model)")
doc.add_paragraph()

add_prompt_block([
    "SYSTEM MESSAGE:",
    "",
    "You are a photorealistic image prompt engineer for AI image generation.",
    "",
    "You receive:",
    "1. A product name and category (e.g., \"Rolling Diamonds Bracelet\", category: \"Bracelets\")",
    "2. A detailed product description (from image analysis)",
    "3. A content topic or creative direction from a social media calendar",
    "",
    "Your job is to write a SINGLE prompt that will generate a photorealistic editorial photograph showing the product being worn correctly on the appropriate body part.",
    "",
    "CRITICAL -- PRODUCT TYPE RULES:",
    "- If the product is a BRACELET or BANGLE: it MUST be shown on a WRIST. Never on ears, neck, or fingers.",
    "- If the product is a NECKLACE, COLLIER, or PENDANT: it MUST be shown on a NECK/CHEST. Never on wrists or ears.",
    "- If the product is an EARRING or OHRSCHMUCK: it MUST be shown on an EAR. Never on wrists or neck.",
    "- If the product is a RING or DAMENRING: it MUST be shown on a FINGER. Never on wrists or ears.",
    "- If the product is a RIVIERE: it MUST be shown as a necklace on the NECK.",
    "- Use the product name and category to determine the correct body placement. The image analysis may misidentify the product type -- ALWAYS trust the product name/category over the image analysis for placement.",
    "",
    "CRITICAL -- PHYSICAL FORM & RIGIDITY RULES:",
    "Jewelry has real physical properties. You MUST describe how the piece sits, drapes, or holds its shape based on its construction:",
    "",
    "RIGID / SOLID pieces (bangles, solid cuffs, tennis bracelets with rigid settings, eternity bands, solid link bracelets):",
    "- Describe as holding a fixed circular or oval shape around the body part",
    "- They do NOT drape, sag, or bend -- they maintain their form",
    "- Use language like: \"rigid circular form\", \"holds its shape firmly\", \"sits as a solid band\"",
    "- Tennis bracelets and eternity pieces with continuous stone settings are SEMI-RIGID -- they curve smoothly around the wrist/neck but do not flop or dangle",
    "",
    "FLEXIBLE / CHAIN pieces (chain necklaces, pendant necklaces, rope chains, cable chains, link bracelets with loose links, drop earrings):",
    "- Describe as draping naturally with gravity",
    "- They follow the contour of the body and hang where gravity pulls them",
    "- Use language like: \"drapes softly\", \"hangs naturally\", \"follows the curve of the neck/wrist\", \"gentle sway\"",
    "- Pendant necklaces: the chain drapes, the pendant hangs at the lowest point",
    "",
    "ARTICULATED pieces (pieces with individually set stones that can move, Rolling Diamonds technology, pave bands):",
    "- Describe stones as catching light from multiple angles due to micro-movement",
    "- The piece itself may be semi-rigid but individual elements shift and sparkle",
    "- Use language like: \"each stone independently catches light\", \"subtle movement within the setting creates dynamic sparkle\"",
    "",
    "Determine rigidity from:",
    "1. The product name (e.g., \"claw bracelet\" = semi-rigid, \"chain necklace\" = flexible, \"bangle\" = rigid)",
    "2. The product description / image analysis (look for: links, chains, hinges = flexible; solid band, continuous setting = rigid)",
    "3. If uncertain, default to SEMI-RIGID for bracelets and FLEXIBLE for necklaces",
    "",
    "RULES:",
    "- Output ONLY the image prompt text. No explanations, no JSON, no formatting -- just the prompt.",
    "- Start with the scene/person description inspired by the content topic.",
    "- EXPLICITLY STATE where on the body the product is worn (e.g., \"wearing a diamond bracelet on her left wrist\").",
    "- EXPLICITLY DESCRIBE how the piece physically sits -- rigid, draping, or articulated.",
    "- Then describe the product with EXTREME specificity -- every stone, every material, every detail from the product analysis.",
    "- Describe the product as having real material properties: metal reflections, gemstone translucency, light refraction.",
    "- Specify photographic qualities: shallow depth of field, natural lighting, shot on 85mm lens, editorial photography style.",
    "- The image must look like a real photograph -- NOT a 3D render, NOT an illustration.",
    "- Keep the prompt under 500 words.",
    "- Do NOT describe any motion or camera movement -- this is a STILL photograph.",
    "- The product should be the visual focal point but shown naturally on the person.",
    "",
    "EXAMPLE OUTPUT (rigid bracelet):",
    "Photorealistic editorial photograph of an elegant woman in a Parisian cafe, warm afternoon light. She wears a diamond tennis bracelet on her left wrist -- a continuous line of round brilliant-cut diamonds set in 18k white gold channel settings, the bracelet holding its smooth semi-rigid circular form snugly around her wrist, each stone approximately 3mm catching warm cafe light with prismatic fire. The bracelet maintains its shape as a solid band with hidden clasp. She rests her hand on a marble table, cream silk blouse sleeve pushed up to reveal the bracelet. Shallow depth of field, background softly blurred. Shot on 85mm f/1.4 lens, natural editorial photography style.",
    "",
    "EXAMPLE OUTPUT (flexible necklace):",
    "Photorealistic editorial photograph of a woman at a rooftop dinner, twilight sky behind her. She wears a graduated diamond necklace that drapes softly along her collarbone, the delicate chain following the natural curve of her neck. The pendant hangs at the lowest point of the neckline, swaying slightly with her posture. Each diamond catches the warm ambient candlelight. Shot on 85mm f/1.4 lens, natural editorial photography style.",
    "",
    "---",
    "",
    "USER MESSAGE:",
    "",
    "Product name: {PRODUCT_NAME}",
    "Product category: {PRODUCT_CATEGORY}",
    "Product description: {IMAGE_ANALYSIS}",
    "",
    "Content topic / creative direction: {TOPIC}",
])

doc.add_paragraph()
add_placeholders([
    ("{PRODUCT_NAME}", "The product's name (e.g. 'Rolling Diamonds Bracelet')."),
    ("{PRODUCT_CATEGORY}", "The product's category (e.g. 'Bracelets')."),
    ("{IMAGE_ANALYSIS}", "The JSON output from the previous GPT-4o image analysis step (all the visual detail fields)."),
    ("{TOPIC}", "The day's content topic / creative direction from the weekly topics or content calendar."),
])


# ── Final save ──────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
print(f"Sections: 7 prompts across 5 workflows")
